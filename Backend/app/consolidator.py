# backend/app/consolidator.py
from docx import Document
import uuid
from pathlib import Path
from typing import List, Dict, Any

OUT_DIR = Path(__file__).resolve().parents[1] / "out"
OUT_DIR.mkdir(exist_ok=True)


def make_proposal(
    requirements: List[Dict[str, Any]],
    applicant_info: Dict[str, Any],
    pricing: Dict[str, Any]
) -> str:
    doc = Document()

    # Header
    doc.add_heading("Proposal", level=1)
    doc.add_paragraph(f"Applicant: {applicant_info.get('name', 'Unknown')}")

    # Requirements
    doc.add_heading("Requirements", level=2)
    for r in requirements:
        text = r.get("text", "")
        doc.add_paragraph(text)

    # Pricing
    doc.add_heading("Pricing", level=2)
    for li in pricing.get("line_items", []):
        sku = li.get("sku", "UNKNOWN")
        qty = li.get("qty", 1)
        amount = li.get("amount", 0)
        doc.add_paragraph(f"{sku} x {qty}: {amount}")

    total = pricing.get("total", 0)
    doc.add_paragraph(f"Total: {total}")

    # Filename
    fname = OUT_DIR / f"proposal_{uuid.uuid4().hex}.docx"

    # FIX: Convert Path → str for Pylance + python-docx
    doc.save(str(fname))

    return str(fname)
